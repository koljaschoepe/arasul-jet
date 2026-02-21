"""
Metadata Extractor for Document Intelligence System
Extracts comprehensive metadata from PDF, DOCX, and Markdown files
"""

import os
import re
import logging
from io import BytesIO
from typing import Dict, Any, Optional, List
from datetime import datetime

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


def extract_metadata(file_data: bytes, filename: str, file_extension: str) -> Dict[str, Any]:
    """
    Extract comprehensive metadata from a document

    Args:
        file_data: Raw file bytes
        filename: Original filename
        file_extension: File extension (e.g., '.pdf')

    Returns:
        Dictionary containing extracted metadata
    """
    metadata = {
        'title': None,
        'author': None,
        'language': 'de',  # Default to German
        'page_count': None,
        'word_count': 0,
        'char_count': 0,
        'creation_date': None,
        'modification_date': None,
        'subject': None,
        'keywords': [],
        'extracted_text_preview': '',
    }

    try:
        ext = file_extension.lower()

        if ext == '.pdf':
            metadata = extract_pdf_metadata(file_data, metadata)
        elif ext == '.docx':
            metadata = extract_docx_metadata(file_data, metadata)
        elif ext in ['.md', '.markdown']:
            metadata = extract_markdown_metadata(file_data, metadata)
        elif ext == '.txt':
            metadata = extract_txt_metadata(file_data, metadata)

        # Fallback title extraction from filename
        if not metadata['title']:
            # Remove extension and clean up filename
            clean_name = os.path.splitext(filename)[0]
            clean_name = re.sub(r'[_-]', ' ', clean_name)
            clean_name = re.sub(r'\s+', ' ', clean_name).strip()
            metadata['title'] = clean_name[:200] if clean_name else filename

        # Detect language from text preview
        if metadata['extracted_text_preview']:
            metadata['language'] = detect_language(metadata['extracted_text_preview'])

    except Exception as e:
        logger.error(f"Error extracting metadata from {filename}: {e}")

    return metadata


def extract_pdf_metadata(file_data: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Extract metadata from PDF file using PyMuPDF"""
    try:
        doc = fitz.open(stream=file_data, filetype="pdf")

        # Page count
        metadata['page_count'] = len(doc)

        # PDF metadata
        pdf_meta = doc.metadata
        if pdf_meta:
            metadata['title'] = pdf_meta.get('title') or None
            metadata['author'] = pdf_meta.get('author') or None
            metadata['subject'] = pdf_meta.get('subject') or None

            # Extract keywords
            keywords = pdf_meta.get('keywords', '')
            if keywords:
                metadata['keywords'] = [k.strip() for k in keywords.split(',')]

            # Dates
            creation = pdf_meta.get('creationDate', '')
            if creation:
                metadata['creation_date'] = parse_pdf_date(creation)

            modification = pdf_meta.get('modDate', '')
            if modification:
                metadata['modification_date'] = parse_pdf_date(modification)

        # Extract text for word/char count and preview
        full_text = []
        for page_num in range(min(10, len(doc))):
            text = doc[page_num].get_text()
            if text:
                full_text.append(text)

        doc.close()

        combined_text = '\n'.join(full_text)
        metadata['char_count'] = len(combined_text)
        metadata['word_count'] = len(combined_text.split())
        metadata['extracted_text_preview'] = combined_text[:2000]

    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {e}")

    return metadata


def extract_docx_metadata(file_data: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Extract metadata from DOCX file"""
    try:
        file_obj = BytesIO(file_data)
        doc = Document(file_obj)

        # Core properties
        if doc.core_properties:
            props = doc.core_properties
            metadata['title'] = props.title
            metadata['author'] = props.author
            metadata['subject'] = props.subject
            metadata['keywords'] = props.keywords.split(',') if props.keywords else []
            metadata['creation_date'] = props.created
            metadata['modification_date'] = props.modified

        # Extract text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Include tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_text.append(row_text)

        combined_text = '\n'.join(full_text)
        metadata['char_count'] = len(combined_text)
        metadata['word_count'] = len(combined_text.split())
        metadata['extracted_text_preview'] = combined_text[:2000]

        # Estimate page count (roughly 500 words per page)
        metadata['page_count'] = max(1, metadata['word_count'] // 500)

    except Exception as e:
        logger.error(f"Error extracting DOCX metadata: {e}")

    return metadata


def extract_markdown_metadata(file_data: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Extract metadata from Markdown file"""
    try:
        text = decode_text(file_data)

        # Extract YAML frontmatter if present
        frontmatter = extract_yaml_frontmatter(text)
        if frontmatter:
            metadata['title'] = frontmatter.get('title')
            metadata['author'] = frontmatter.get('author')
            metadata['keywords'] = frontmatter.get('tags', [])
            if isinstance(metadata['keywords'], str):
                metadata['keywords'] = [k.strip() for k in metadata['keywords'].split(',')]

        # If no title in frontmatter, extract from first heading
        if not metadata['title']:
            title_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
            if title_match:
                metadata['title'] = title_match.group(1).strip()

        metadata['char_count'] = len(text)
        metadata['word_count'] = len(text.split())
        metadata['extracted_text_preview'] = text[:2000]

        # Estimate page count
        metadata['page_count'] = max(1, metadata['word_count'] // 500)

    except Exception as e:
        logger.error(f"Error extracting Markdown metadata: {e}")

    return metadata


def extract_txt_metadata(file_data: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Extract metadata from plain text file"""
    try:
        text = decode_text(file_data)

        metadata['char_count'] = len(text)
        metadata['word_count'] = len(text.split())
        metadata['extracted_text_preview'] = text[:2000]

        # Estimate page count
        metadata['page_count'] = max(1, metadata['word_count'] // 500)

    except Exception as e:
        logger.error(f"Error extracting TXT metadata: {e}")

    return metadata


def extract_yaml_frontmatter(text: str) -> Optional[Dict[str, Any]]:
    """Extract YAML frontmatter from markdown"""
    try:
        if text.startswith('---'):
            end_match = re.search(r'\n---\s*\n', text[3:])
            if end_match:
                yaml_content = text[3:end_match.start() + 3]
                # Simple YAML parsing (key: value format)
                frontmatter = {}
                for line in yaml_content.split('\n'):
                    if ':' in line:
                        key, value = line.split(':', 1)
                        frontmatter[key.strip()] = value.strip().strip('"\'')
                return frontmatter
    except Exception:
        pass
    return None


def decode_text(file_data: bytes) -> str:
    """Decode bytes to text with multiple encoding attempts"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    for encoding in encodings:
        try:
            return file_data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_data.decode('utf-8', errors='ignore')


def parse_pdf_date(date_str: str) -> Optional[datetime]:
    """Parse PDF date format (D:YYYYMMDDHHmmSS)"""
    try:
        if date_str.startswith('D:'):
            date_str = date_str[2:]
        # Extract just the date/time portion
        date_str = re.sub(r'[^\d]', '', date_str[:14])
        if len(date_str) >= 8:
            return datetime.strptime(date_str[:14].ljust(14, '0'), '%Y%m%d%H%M%S')
    except Exception:
        pass
    return None


def detect_language(text: str) -> str:
    """
    Simple language detection based on common words
    Returns ISO 639-1 code
    """
    text_lower = text.lower()

    # German indicators
    german_words = ['und', 'der', 'die', 'das', 'ist', 'ein', 'eine', 'für', 'mit',
                    'auf', 'werden', 'wird', 'kann', 'auch', 'nicht', 'sind', 'haben']
    german_count = sum(1 for word in german_words if f' {word} ' in f' {text_lower} ')

    # English indicators
    english_words = ['the', 'and', 'is', 'are', 'for', 'with', 'can', 'also',
                     'not', 'have', 'this', 'that', 'from', 'will', 'would']
    english_count = sum(1 for word in english_words if f' {word} ' in f' {text_lower} ')

    if german_count > english_count:
        return 'de'
    elif english_count > german_count:
        return 'en'
    else:
        # Default to German for this system
        return 'de'


def extract_key_topics(text: str, max_topics: int = 10) -> List[str]:
    """
    Extract key topics/keywords from text using simple TF analysis
    For more sophisticated extraction, use the LLM categorizer
    """
    # Remove common stopwords
    stopwords = {
        'de': {'und', 'der', 'die', 'das', 'ist', 'ein', 'eine', 'für', 'mit', 'auf',
               'werden', 'wird', 'kann', 'auch', 'nicht', 'sind', 'haben', 'oder', 'von',
               'zu', 'an', 'bei', 'nach', 'aus', 'wenn', 'als', 'wie', 'so', 'es'},
        'en': {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
               'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'been',
               'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
               'could', 'should', 'may', 'might', 'must', 'shall', 'can', 'this', 'that'}
    }

    # Tokenize and count words
    words = re.findall(r'\b[a-zäöüß]{4,}\b', text.lower())

    # Filter stopwords
    all_stopwords = stopwords['de'] | stopwords['en']
    words = [w for w in words if w not in all_stopwords]

    # Count frequency
    word_freq = {}
    for word in words:
        word_freq[word] = word_freq.get(word, 0) + 1

    # Get top words
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    topics = [word for word, _ in sorted_words[:max_topics]]

    return topics
