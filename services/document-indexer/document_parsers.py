"""
Parser je Dateiformat und der gemeinsame Einstieg `parse_document`.

Unterstuetzt PDF (PyMuPDF fuer Text, pdfplumber fuer Tabellen, OCR-Rueckfall
fuer Bild-PDFs), DOCX, TXT, Markdown, HTML/XML, CSV/JSON/Log, YAML-Tabellen
und Bilder (OCR).

Seit Phase B4 (26.08.2026) liegt hier auch das, was vorher in
`document_processor.py` stand und die Extraktion braucht: `PARSERS`,
`strip_nul` und `parse_document`. Der Rest jenes Moduls (Chunking,
Textlayer, Anreicherung) ist mit dem Hintergrund-Indexer gefallen.
"""

import gc
import logging
import os
from io import BytesIO
from typing import IO, Generator, Optional

import fitz  # PyMuPDF: Textextraktion mit Layouterhalt
import pdfplumber  # Tabellen aus PDFs
from docx import Document
import yaml

# OCR service for scanned document support
try:
    from ocr_service import parse_pdf_with_ocr_fallback
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)


def _format_tables_from_page(page) -> str:
    """
    Format all tables found on an already-open pdfplumber page object.

    Separated from PDF opening so callers processing many pages can open the
    document with ``pdfplumber.open`` ONCE and reuse it, instead of re-parsing
    the whole multi-MB buffer per page (O(pages) full reparses).

    Args:
        page: An open ``pdfplumber`` page object

    Returns:
        Pipe-delimited text representation of all tables found on the page,
        or empty string if no tables found
    """
    try:
        tables = page.extract_tables()

        if not tables:
            return ""

        table_texts = []
        for table_idx, table in enumerate(tables):
            if not table:
                continue

            rows_text = []
            for row in table:
                # Replace None values with empty string
                cleaned_row = [
                    str(cell).strip() if cell is not None else ""
                    for cell in row
                ]
                rows_text.append(" | ".join(cleaned_row))

            if rows_text:
                table_texts.append("\n".join(rows_text))

        return "\n\n".join(table_texts)

    except Exception as e:
        logger.debug(f"Table extraction failed: {e}")
        return ""


def parse_pdf(file_obj: IO[bytes], use_ocr: bool = True) -> str:
    """
    Parse PDF file and extract text using PyMuPDF (fitz) for high-quality
    text extraction with layout preservation, plus pdfplumber for table extraction.
    Automatically falls back to OCR for scanned documents if OCR is available.

    Args:
        file_obj: File object containing PDF data
        use_ocr: Whether to attempt OCR for scanned PDFs (default: True)

    Returns:
        Extracted text from PDF
    """
    try:
        # Try OCR-enabled parsing if available and enabled
        if use_ocr and OCR_AVAILABLE:
            text, used_ocr = parse_pdf_with_ocr_fallback(file_obj)
            if used_ocr:
                logger.info("PDF parsed using OCR")
            return text

        # Standard extraction using PyMuPDF + pdfplumber
        file_obj.seek(0)
        pdf_bytes = file_obj.read()

        # Open with fitz for text extraction
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text_parts = []

        # Open pdfplumber ONCE for the whole document and reuse it per page,
        # instead of re-parsing the entire buffer on every page.
        with pdfplumber.open(BytesIO(pdf_bytes)) as plumber_pdf:
            for page_num in range(len(doc)):
                page = doc[page_num]

                # Extract text with layout preservation
                # Using "text" sort mode for natural reading order
                text = page.get_text("text")

                page_content = ""
                if text and text.strip():
                    page_content = text.strip()

                # Extract tables using pdfplumber and append to page text
                table_text = ""
                if page_num < len(plumber_pdf.pages):
                    table_text = _format_tables_from_page(plumber_pdf.pages[page_num])
                if table_text:
                    if page_content:
                        page_content += "\n\n[Table]\n" + table_text
                    else:
                        page_content = "[Table]\n" + table_text

                if page_content:
                    text_parts.append(page_content)

        doc.close()

        full_text = "\n\n".join(text_parts)
        return full_text.strip()

    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise


def parse_image(file_obj: IO[bytes]) -> str:
    """
    Parse image file using OCR to extract text.

    Args:
        file_obj: File object containing image data (PNG, JPEG, etc.)

    Returns:
        Extracted text from image, or empty string if OCR not available
    """
    if not OCR_AVAILABLE:
        logger.warning("OCR not available, cannot extract text from image")
        return ""

    try:
        from ocr_service import ocr_image

        file_obj.seek(0)
        image_data = file_obj.read()

        result = ocr_image(image_data)

        if result.success:
            logger.info(f"Image OCR successful: {len(result.text)} chars extracted")
            return result.text
        else:
            logger.warning(f"Image OCR failed: {result.error}")
            return ""

    except Exception as e:
        logger.error(f"Error parsing image with OCR: {e}")
        return ""


def parse_pdf_streaming(file_obj: IO[bytes], gc_interval: int = 10) -> Generator[str, None, None]:
    """
    MEDIUM-PRIORITY-FIX 3.3: Memory-efficient streaming PDF parser

    Parse PDF file page by page using a generator to reduce memory usage.
    Uses PyMuPDF (fitz) for high-quality text extraction and pdfplumber
    for table extraction.
    Useful for large PDFs (100+ pages) where loading all text at once
    would cause memory spikes.

    Args:
        file_obj: File object containing PDF data
        gc_interval: Run garbage collection every N pages (default: 10)

    Yields:
        Text content from each page

    Example:
        for page_text in parse_pdf_streaming(file_obj):
            # Process page_text chunk by chunk
            chunks = chunk_text(page_text)
            for chunk in chunks:
                process_chunk(chunk)
    """
    try:
        file_obj.seek(0)
        pdf_bytes = file_obj.read()

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc)
        logger.info(f"Starting streaming PDF parse: {total_pages} pages")

        # Open pdfplumber ONCE for the whole document and reuse it per page,
        # instead of re-parsing the entire buffer on every page (O(pages)
        # full reparses of a multi-MB buffer).
        with pdfplumber.open(BytesIO(pdf_bytes)) as plumber_pdf:
            for page_num in range(total_pages):
                try:
                    page = doc[page_num]
                    text = page.get_text("text")

                    page_content = ""
                    if text and text.strip():
                        page_content = text.strip()

                    # Extract tables using pdfplumber
                    table_text = ""
                    if page_num < len(plumber_pdf.pages):
                        table_text = _format_tables_from_page(plumber_pdf.pages[page_num])
                    if table_text:
                        if page_content:
                            page_content += "\n\n[Table]\n" + table_text
                        else:
                            page_content = "[Table]\n" + table_text

                    if page_content:
                        yield page_content

                    # Periodic garbage collection for very large PDFs
                    if gc_interval > 0 and (page_num + 1) % gc_interval == 0:
                        gc.collect()
                        logger.debug(f"Processed {page_num + 1}/{total_pages} pages (GC triggered)")

                except Exception as page_error:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {page_error}")
                    # Continue with next page instead of failing entire document
                    continue

        doc.close()
        logger.info(f"Completed streaming PDF parse: {total_pages} pages processed")

    except Exception as e:
        logger.error(f"Error in streaming PDF parse: {e}")
        raise


def get_pdf_page_count(file_obj: IO[bytes]) -> int:
    """
    Get the number of pages in a PDF without extracting text.
    Useful for progress tracking and deciding whether to use streaming parser.

    Args:
        file_obj: File object containing PDF data

    Returns:
        Number of pages in the PDF
    """
    try:
        file_obj.seek(0)
        pdf_bytes = file_obj.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        count = len(doc)
        doc.close()
        file_obj.seek(0)  # Reset for subsequent reads
        return count
    except Exception as e:
        logger.error(f"Error getting PDF page count: {e}")
        return 0


def parse_docx(file_obj: IO[bytes]) -> str:
    """
    Parse DOCX file and extract text

    Args:
        file_obj: File object containing DOCX data

    Returns:
        Extracted text from DOCX
    """
    try:
        doc = Document(file_obj)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)
        return full_text.strip()

    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        raise


def parse_txt(file_obj: IO[bytes]) -> str:
    """
    Parse plain text file

    Args:
        file_obj: File object containing text data

    Returns:
        Text content
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                file_obj.seek(0)
                text = file_obj.read().decode(encoding)
                return text.strip()
            except UnicodeDecodeError:
                continue

        # If all encodings fail, use utf-8 with errors='ignore'
        file_obj.seek(0)
        text = file_obj.read().decode('utf-8', errors='ignore')
        return text.strip()

    except Exception as e:
        logger.error(f"Error parsing TXT: {e}")
        raise


def parse_markdown(file_obj: IO[bytes]) -> str:
    """
    Parse Markdown file and extract text
    Converts markdown to plain text by rendering and stripping HTML

    Args:
        file_obj: File object containing markdown data

    Returns:
        Plain text extracted from markdown
    """
    try:
        # Read markdown content
        md_text = parse_txt(file_obj)

        # For RAG purposes, we keep the markdown formatting
        # as it preserves structure (headers, lists, etc.)
        # which can be useful for context
        return md_text

    except Exception as e:
        logger.error(f"Error parsing Markdown: {e}")
        raise


def parse_html(file_obj: IO[bytes]) -> str:
    """
    Parse HTML/XML file and extract readable text.

    Strips markup (including <script>/<style> contents), unescapes
    entities, and normalizes whitespace so the result chunks cleanly.

    Args:
        file_obj: File object containing HTML/XML data

    Returns:
        Plain text extracted from the markup
    """
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        SKIP_TAGS = {'script', 'style', 'noscript', 'template'}
        BLOCK_TAGS = {
            'p', 'div', 'section', 'article', 'header', 'footer', 'main',
            'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'tr', 'br', 'hr',
            'table', 'ul', 'ol', 'blockquote', 'pre', 'td', 'th'
        }

        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.parts = []
            self._skip_depth = 0

        def handle_starttag(self, tag, attrs):
            if tag in self.SKIP_TAGS:
                self._skip_depth += 1
            elif tag in self.BLOCK_TAGS:
                self.parts.append('\n')

        def handle_endtag(self, tag):
            if tag in self.SKIP_TAGS and self._skip_depth > 0:
                self._skip_depth -= 1
            elif tag in self.BLOCK_TAGS:
                self.parts.append('\n')

        def handle_data(self, data):
            if not self._skip_depth and data.strip():
                self.parts.append(data)

    try:
        raw = parse_txt(file_obj)
        extractor = _TextExtractor()
        extractor.feed(raw)
        extractor.close()
        text = ''.join(extractor.parts)
        lines = [' '.join(line.split()) for line in text.splitlines()]
        return '\n'.join(line for line in lines if line).strip()

    except Exception as e:
        logger.error(f"Error parsing HTML: {e}")
        raise


def parse_yaml_table(file_obj: IO[bytes]) -> str:
    """
    Parse YAML table file and convert to searchable text for RAG.
    Extracts table metadata, column names, and all row data.

    Args:
        file_obj: File object containing YAML table data

    Returns:
        Formatted text representation of the table for indexing

    YAML Table Format:
        _meta:
            name: "Table Name"
            description: "Description"
        columns:
            - slug: "col1"
              name: "Column 1"
              type: "text"
        rows:
            - col1: "value1"
              col2: "value2"
    """
    try:
        # Read YAML content
        content = parse_txt(file_obj)
        data = yaml.safe_load(content)

        text_parts = []

        # Nur ein Mapping kann das Arasul-Tabellenformat tragen. Alles
        # andere (Top-Level-Liste, Skalar, None) faellt direkt in den
        # Text-Fallback unten — vorher lief eine Liste hier in ein
        # AttributeError auf `data.get`, und die Datei galt als
        # „nicht parsebar".
        if isinstance(data, dict):
            # Extract metadata
            meta = data.get('_meta', {})
            if meta.get('name'):
                text_parts.append(f"Tabelle: {meta['name']}")
            if meta.get('description'):
                text_parts.append(f"Beschreibung: {meta['description']}")

            # Extract column information
            columns = data.get('columns', [])
            if columns:
                col_names = [c.get('name', c.get('slug', '')) for c in columns]
                text_parts.append(f"Spalten: {', '.join(col_names)}")

                # Create a mapping of slug to name for row formatting
                slug_to_name = {c.get('slug', ''): c.get('name', c.get('slug', '')) for c in columns}

            # Extract row data
            rows = data.get('rows', [])
            if rows:
                text_parts.append(f"\nDaten ({len(rows)} Einträge):\n")

                for i, row in enumerate(rows):
                    # Skip internal fields
                    row_values = []
                    for key, value in row.items():
                        if key.startswith('_'):
                            continue
                        if value is not None and value != '':
                            # Use column name if available, otherwise use key
                            col_name = slug_to_name.get(key, key) if columns else key
                            row_values.append(f"{col_name}: {value}")

                    if row_values:
                        text_parts.append(' | '.join(row_values))

                    # Limit to first 500 rows for very large tables
                    if i >= 499:
                        text_parts.append(f"... und {len(rows) - 500} weitere Einträge")
                        break

        if text_parts:
            return '\n'.join(text_parts)

        # Kein Arasul-Tabellen-YAML (kein _meta/columns/rows), sondern
        # irgendein anderes: OpenAPI-Spec, CI-Konfiguration,
        # docker-compose. Frueher fiel so etwas hier als LEERER String
        # heraus und galt als „geparst, aber kein Text" — die Datei landete
        # auf 'stored' und war inhaltlich nicht auffindbar, ohne dass
        # irgendwo ein Hinweis auftauchte. Der Rohtext ist fuer die Suche
        # allemal besser als nichts und behaelt die Kommentare, die
        # `safe_load` wegwirft.
        if content.strip():
            logger.info(
                "YAML ohne Tabellenstruktur (_meta/columns/rows) — als Text indexiert"
            )
            return content
        return ""

    except yaml.YAMLError as e:
        logger.error(f"Error parsing YAML: {e}")
        # Return raw content if YAML parsing fails
        return parse_txt(file_obj)
    except Exception as e:
        logger.error(f"Error parsing YAML table: {e}")
        raise


# ---------------------------------------------------------------------------
# Gemeinsamer Einstieg
# ---------------------------------------------------------------------------

STREAMING_PDF_THRESHOLD = 50  # ab so vielen Seiten seitenweise parsen


def parse_pdf_smart(file_obj: IO[bytes]) -> str:
    """Grosse PDFs (> STREAMING_PDF_THRESHOLD Seiten) seitenweise, um Speicher zu sparen."""
    page_count = get_pdf_page_count(file_obj)
    if page_count > STREAMING_PDF_THRESHOLD:
        logger.info(f"Large PDF ({page_count} pages), using streaming parser")
        return "\n\n".join(parse_pdf_streaming(file_obj))
    return parse_pdf(file_obj)


# Parser je Dateiendung
PARSERS = {
    '.pdf': parse_pdf_smart,
    '.txt': parse_txt,
    '.md': parse_markdown,
    '.markdown': parse_markdown,
    '.docx': parse_docx,
    '.yaml': parse_yaml_table,
    '.yml': parse_yaml_table,
    # Markup -> Klartext
    '.html': parse_html,
    '.htm': parse_html,
    '.xml': parse_html,
    # Textbasierte Datenformate
    '.csv': parse_txt,
    '.json': parse_txt,
    '.log': parse_txt,
    # Bilder (OCR)
    '.png': parse_image,
    '.jpg': parse_image,
    '.jpeg': parse_image,
    '.tiff': parse_image,
    '.tif': parse_image,
    '.bmp': parse_image,
    '.webp': parse_image,
}


def strip_nul(text: str) -> str:
    """
    Entfernt NUL-Bytes (0x00) aus extrahiertem Text.

    Manche PDFs liefern beim Extrahieren NUL-Bytes mit (fehlerhafte
    Font-/Encoding-Tabellen). Frueher scheiterte daran der Schreibversuch in
    Postgres; heute schreibt der Indexer nichts mehr, aber das Backend legt
    den Text weiter in seiner Datenbank ab. Hier laufen alle Parser zusammen,
    also wird hier bereinigt.
    """
    if not text:
        return text
    if '\x00' not in text:
        return text
    bereinigt = text.replace('\x00', '')
    logger.warning(
        f"{len(text) - len(bereinigt)} NUL-Byte(s) aus extrahiertem Text entfernt"
    )
    return bereinigt


def parse_document(data: bytes, filename: str) -> Optional[str]:
    """
    Dokument parsen und Text liefern.

    Rueckgabe-Vertrag (der Aufrufer unterscheidet die beiden Faelle):
      None  -> nicht parsebar (unbekannter Typ oder Parser-Fehler)
      ''    -> sauber geparst, aber ohne Textinhalt (z. B. Logo ohne Schrift)
    """
    _, ext = os.path.splitext(filename.lower())

    if ext not in PARSERS:
        logger.warning(f"Unsupported file type: {ext}")
        return None

    try:
        parser = PARSERS[ext]
        text = strip_nul(parser(BytesIO(data)))
        logger.debug(f"Parsed {filename}: {len(text)} characters")
        return text
    except Exception as e:  # noqa: BLE001 - Vertrag: None statt Ausnahme
        logger.error(f"Parse error for {filename}: {e}")
        return None
